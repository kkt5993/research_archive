#!/usr/bin/env python3
"""AI 에이전트 생태계 20종목 — Word 발간본 생성."""
import numpy as np, pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

D = pd.read_csv('agent20_holdings.csv').set_index('ticker')
G = pd.read_csv('agent20_agg.csv', index_col=0).squeeze()
U = pd.read_csv('sp500_universe.csv').set_index('ticker')
held = D[D.weight.notna() & (D.weight > 0)]
NAVY, RED, GREY = RGBColor(0x1F, 0x4E, 0x79), RGBColor(0xC0, 0, 0), RGBColor(0x60, 0x60, 0x60)

doc = Document()
st = doc.styles['Normal']; st.font.name = '맑은 고딕'; st.font.size = Pt(9.5)
st._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(1.8); s.left_margin = s.right_margin = Cm(1.8)

def H(t, size=15, color=NAVY, space_before=10, rule=False):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    if rule:
        pr = p._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '8'); b.set(qn('w:color'), '1F4E79')
        bd.append(b); pr.append(bd)
    return p

def P(t, size=9.5, italic=False, color=None, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t); r.font.size = Pt(size); r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc); tcPr.append(sh)

def table(headers, rows, widths=None, size=8.5, hdr_fill='1F4E79', align_right=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(str(h)); r.bold = True; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, hdr_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            txt = str(v); bold = txt.startswith('**') and txt.endswith('**')
            if bold: txt = txt.strip('*')
            r = cells[i].paragraphs[0].add_run(txt); r.font.size = Pt(size); r.bold = bold
            if txt.startswith('-') or txt.startswith('−'): r.font.color.rgb = RED
            cells[i].paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.RIGHT
                if (align_right and i in align_right) else
                (WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT))
        if ri % 2 == 1:
            for c in cells: shade(c, 'F2F6FA')
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    return t

def f(v, sp='{:.1f}', pct=False, dash='—'):
    if v is None or (isinstance(v, float) and (np.isnan(v) or not np.isfinite(v))): return dash
    return sp.format(v * 100 if pct else v)

# ───────── 표지
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AI 에이전트 생태계 모델 포트폴리오'); r.bold = True; r.font.size = Pt(21); r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('20선 · 메모리·스토리지·소부장·클라우드 집중형'); r.font.size = Pt(12); r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026년 8월 26일 · BM: S&P 500 · 상상인증권 투자전략팀'); r.font.size = Pt(9.5); r.font.color.rgb = GREY

H('설계 방침', 13, rule=True)
table(['항목', '내용'], [
    ['CPU', '**전량 미보유** — INTC·AMD·ARM 0%. BM 대비 −1.63%p'],
    ['DRAM·NAND', '14.5% (BM 1.68%) — MU·SNDK'],
    ['스토리지', '10.0% (BM 0.51%) — WDC·STX·NTAP'],
    ['반도체 소부장', '20.0% (BM 1.36%) — AMAT·LRCX·KLAC·ASML·ENTG'],
    ['클라우드(대형)', '26.0% (BM 15.01%) — MSFT·AMZN·ORCL·GOOGL'],
    ['네오클라우드', '11.0% (BM 0%) — CRWV·NBIS·IREN'],
    ['AI 인프라', '16.0% — NVDA·AVGO'],
    ['GOOGL', '**언더웨이트** 3.0% (BM 5.75%, −2.75%p)'],
    ['TSLA', '**언더웨이트** 0% 보유 (BM 1.86%, −1.86%p)'],
    ['SPCX', '**중립** 2.5% — 시총 비중 2.46% 수준'],
], widths=[3.4, 13.6], size=9)

# ───────── 요약 지표
H('요약 지표', 13, rule=True)
table(['구분', '지표', '값', '지표', '값'], [
    ['수익', '기대수익률(컨센 목표주가)', f'**{f(G.target_upside,"{:+.1f}%",pct=True)}**', '가중 forward P/E', f'{f(G.fwd_pe)}'],
    ['', '컨센 EPS 성장(2년 CAGR)', f'{f(G.growth_2y,"{:+.1f}%",pct=True)}', '가중 PSR', f'{f(G.psr)}'],
    ['수익성', '영업이익률', f'{f(G.opm,"{:.1f}%",pct=True)}', '매출총이익률', f'{f(G.gpm,"{:.1f}%",pct=True)}'],
    ['', 'ROE', f'{f(G.roe,"{:.1f}%",pct=True)}', '배당수익률', f'{f(G.div_yield,"{:.2f}%")}'],
    ['위험', '추적오차(TE)', f'**{f(G.te,"{:.1f}%",pct=True)}**', '베타', f'**{f(G.beta,"{:.2f}")}**'],
    ['', '연율 변동성', f'{f(G.vol,"{:.1f}%",pct=True)}', '하방편차', f'{f(G.dsd,"{:.1f}%",pct=True)}'],
    ['', '최대낙폭(3년)', f'{f(G.mdd,"{:.1f}%",pct=True)}', 'BM 최대낙폭', f'{f(G.mdd_b,"{:.1f}%",pct=True)}'],
    ['', 'VaR 95% / 99%(일간)', f'{f(G.var95,"{:.2f}%",pct=True)} / {f(G.var99,"{:.2f}%",pct=True)}', '최악 1일 / 21일', f'{f(G.worst_day,"{:.1f}%",pct=True)} / {f(G.worst_21,"{:.1f}%",pct=True)}'],
    ['', '상승 / 하락 캡처', f'{f(G.up_cap,"{:.2f}")} / {f(G.dn_cap,"{:.2f}")}', '상승 / 하락 베타', f'{f(G.beta_up,"{:.2f}")} / **{f(G.beta_dn,"{:.2f}")}**'],
    ['구조', '종목 간 평균 상관', f'**{f(G.avg_corr,"{:.3f}")}**', '분산 효과', f'{f(G.diversification,"{:.1f}%",pct=True)}'],
    ['', 'TOP10 비중', f'{f(G.top10,"{:.1f}%")}', '유효종목수', f'{f(1/G.hhi,"{:.1f}")}'],
    ['', 'BM 밖 비중', f'**{f(G.bm_out,"{:.1f}%")}**', '순부채/EBITDA', f'{f(G.nd_ebitda,"{:.2f}배")}'],
    ['', '가중 애널 커버리지', f'{f(G.n_analysts,"{:.0f}명")}', '가중 일평균거래대금', f'${f(G.adv_wavg/1e9,"{:.1f}B")}'],
], widths=[1.6, 5.0, 3.0, 4.4, 3.0], size=8.5, align_right=[2, 4])

doc.add_page_break()

# ───────── 차트
H('비중 구성', 13, rule=True)
doc.add_picture('agent20_charts.png', width=Cm(17.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
P('테마별 비중(좌) · 종목별 비중(중) · BM 대비 액티브 비중(우). * 표시는 미보유 종목으로, '
  '0% 보유가 곧 최대 언더웨이트다.', 8, italic=True, color=GREY)

# ───────── 보유 종목
H('보유 종목 20선', 13, rule=True)
rows = []
for t, x in held.sort_values('weight', ascending=False).iterrows():
    rows.append([t, x['name'][:18], x.group, f'{x.weight:.1f}',
                 f'{x.active:+.2f}' if x.bm > 0 else f'{x.active:+.2f}*',
                 f(x.price, '{:,.2f}'), f(x.fwd_pe), f(x.growth_2y, '{:+.0f}%', pct=True),
                 f(x.target_upside, '{:+.0f}%', pct=True), f(x.opm, '{:.0f}%', pct=True),
                 f(x.n_analysts, '{:.0f}')])
table(['티커', '종목명', '테마', '비중', '액티브', '주가', 'FwdPE', '성장', '목표여력', '영업M', '커버'],
      rows, widths=[1.3, 3.0, 2.3, 1.1, 1.3, 1.7, 1.2, 1.2, 1.4, 1.1, 1.0], size=8,
      align_right=[3, 4, 5, 6, 7, 8, 9, 10])
P('* BM(S&P 500) 미편입 종목 — 액티브는 보유 비중 전액이다. 성장은 직전 실적 대비 차년도 '
  '컨센서스까지의 2년 CAGR, 목표여력은 애널리스트 목표주가 중앙값 대비 상승여력.', 8, italic=True, color=GREY)

doc.add_page_break()

# ───────── 테마별 논거
H('테마별 논거', 13, rule=True)

H('1. 메모리·스토리지 24.5% — 비중은 올리되 사이클 정점을 전제로 잡는다', 11, RGBColor(0,0,0), 8)
table(['종목', '비중', 'Fwd P/E', '영업이익률', '매출성장', '판단'], [
    ['MU', '9.0', f"{f(D.fwd_pe.get('MU'))}", f"{f(D.opm.get('MU'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('MU'),'{:+.0f}%',pct=True)}", 'HBM 주도, 최대 비중'],
    ['SNDK', '5.5', f"{f(D.fwd_pe.get('SNDK'))}", f"{f(D.opm.get('SNDK'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('SNDK'),'{:+.0f}%',pct=True)}", 'NAND 순수 노출'],
    ['WDC', '4.0', f"{f(D.fwd_pe.get('WDC'))}", f"{f(D.opm.get('WDC'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('WDC'),'{:+.0f}%',pct=True)}", 'HDD·니어라인'],
    ['STX', '4.0', f"{f(D.fwd_pe.get('STX'))}", f"{f(D.opm.get('STX'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('STX'),'{:+.0f}%',pct=True)}", 'HDD 과점'],
    ['NTAP', '2.0', f"{f(D.fwd_pe.get('NTAP'))}", f"{f(D.opm.get('NTAP'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('NTAP'),'{:+.0f}%',pct=True)}", '데이터 관리 소프트'],
], widths=[1.6, 1.4, 1.8, 2.2, 2.0, 8.0], size=8.5, align_right=[1,2,3,4])
P('에이전트는 추론을 반복하며 컨텍스트를 저장한다 — 연산보다 메모리와 스토리지를 먼저 소모하는 '
  '수요 구조다. 이것이 이 축을 24.5%(BM 2.19%)까지 올린 근거다.', 9)
P('다만 MU·SNDK의 영업이익률 78~80%는 역사적 극단이며, forward P/E 5~6배는 싸서가 아니라 '
  '이익이 정점이라 낮게 보이는 배수다. 사이클이 꺾이면 분모가 먼저 무너진다. '
  '이 두 종목의 비중 14.5%는 그 위험을 감수한 값이며, 게이트를 반드시 병행해야 한다.', 9)

H('2. 소부장 20.0% — 사이클 위에서 한 단계 완충된 자리', 11, RGBColor(0,0,0), 8)
P('메모리 업체가 capex를 집행하면 장비·소재는 그 지출을 매출로 받는다. 메모리 가격이 아니라 '
  '메모리 투자에 연동되므로 사이클 진폭이 한 단계 완만하다. AMAT·LRCX·KLAC는 NAND 적층과 '
  'HBM 패키징에 직접 노출되고, ASML은 EUV 독점, ENTG는 소재 축이다.', 9)

H('3. 클라우드(대형) 26.0% — 에이전트가 실행되는 장소', 11, RGBColor(0,0,0), 8)
P('에이전트는 결국 하이퍼스케일러 위에서 돈다. MSFT 10.0%·AMZN 8.0%·ORCL 5.0%로 담되, '
  f'GOOGL은 3.0%로 언더웨이트한다(BM 5.75%, {f(-2.75,"{:+.2f}")}%p).', 9)
P('GOOGL 언더웨이트의 근거는 검색 잠식이 아니다 — 점유율 91.4%, 검색 광고 매출 Q1 +19%·Q2 +17%로 '
  '그 논거는 자료가 반박한다. 근거는 현금흐름이다: 2년간 영업현금흐름 +62%인데 잉여현금흐름은 +5.5%에 '
  '그쳤고, 2026년 capex 가이던스는 $195~205B로 전년의 2.2배다. 2026년 6월 주식·의무전환우선주 발행으로 '
  '순수취 $49.6B를 조달하고 최대 $40B ATM 프로그램을 체결했다(SEC 8-K 원문 확인). '
  '현금흐름으로 감당되는 투자라면 하지 않을 조달이다.', 9)

H('4. 네오클라우드 11.0% — 비중은 올리되 성격은 옵션이다', 11, RGBColor(0,0,0), 8)
table(['종목', '비중', 'Fwd P/E', '영업이익률', '매출성장', '베타(3년)'], [
    ['CRWV', '4.0', f"{f(D.fwd_pe.get('CRWV'))}", f"{f(D.opm.get('CRWV'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('CRWV'),'{:+.0f}%',pct=True)}", f"{f(D.beta_spy.get('CRWV'),'{:.2f}')}"],
    ['NBIS', '4.0', f"{f(D.fwd_pe.get('NBIS'))}", f"{f(D.opm.get('NBIS'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('NBIS'),'{:+.0f}%',pct=True)}", f"{f(D.beta_spy.get('NBIS'),'{:.2f}')}"],
    ['IREN', '3.0', f"{f(D.fwd_pe.get('IREN'))}", f"{f(D.opm.get('IREN'),'{:.0f}%',pct=True)}", f"{f(D.rev_g.get('IREN'),'{:+.0f}%',pct=True)}", f"{f(D.beta_spy.get('IREN'),'{:.2f}')}"],
], widths=[1.6, 1.4, 1.8, 2.2, 2.0, 2.0], size=8.5, align_right=[1,2,3,4,5])
P('셋 다 영업적자이고 베타는 2~4배다. 밸류에이션으로 정당화되지 않으며 정당화하려 하지도 않는다 — '
  '이 11%는 컨빅션이 아니라 옵션이고, 그렇게 부르는 편이 정직하다. 손실 한도를 미리 정하고 '
  '차환·고객 집중도를 월 단위로 점검하는 것이 전제다.', 9)

H('5. CPU 0% · TSLA 0% — 미보유가 곧 최대 언더웨이트', 11, RGBColor(0,0,0), 8)
P(f'INTC·AMD를 전량 배제해 BM 대비 {f(G.cpu_uw,"{:.2f}")}%p, TSLA 배제로 {f(G.tsla_uw,"{:.2f}")}%p를 '
  '언더웨이트한다. 에이전트 워크로드에서 범용 CPU는 병목이 아니며, 추론 확산은 가속기·메모리·'
  '네트워크로 지출을 옮긴다. SPCX는 시총 비중(2.46%) 수준인 2.5%로 중립을 유지한다.', 9)

doc.add_page_break()

# ───────── 위험
H('위험 관리', 13, rule=True)
P(f'이 포트는 추적오차 {f(G.te,"{:.1f}%",pct=True)}, 베타 {f(G.beta,"{:.2f}")}, 연율 변동성 '
  f'{f(G.vol,"{:.1f}%",pct=True)}로 벤치마크 대비 매우 공격적이다. 참고로 동일 팀의 코어 포트폴리오'
  '(MP v5.1, 30종목)는 TE 11.9%·베타 1.37이다 — 이 포트는 그 2.6배의 액티브 위험을 진다.', 9)
table(['게이트', '트리거', '행동'], [
    ['메모리 사이클', '**MU·SNDK 영업이익률 60% 하회** 또는 DRAM 현물가 2개월 연속 하락', '메모리 축 1차 축소(14.5%→8%)'],
    ['CXMT·공급', '중국 DRAM 증설 가속 / MU HBM 점유 정체(월간)', 'MU 축소'],
    ['네오클라우드', '고객 집중도 악화 · 차환 실패 · 계약 취소', '**즉시 전량**'],
    ['소부장', '메모리 3사 capex 가이던스 하향', '소부장 축소'],
    ['금리', '인하 기대 훼손', '고멀티플부터: 네오클라우드 → 소부장'],
    ['GOOGL', 'capex 가이던스 하향 또는 FCF 마진 20% 회복', '언더웨이트 해제'],
    ['집중도', '개별 10% · BM외 20%', '상한 도달 시 교체 편입만'],
], widths=[2.6, 8.6, 5.8], size=8.5)

H('한계 — 감추지 않는 것', 12, rule=True)
P(f'1. BM 밖 비중이 {f(G.bm_out,"{:.1f}%")}다(ASML·ENTG·CRWV·NBIS·IREN·SPCX). 벤치마크 상대평가의 '
  '의미가 그만큼 줄어든다.', 9, after=2)
P(f'2. 종목 간 평균 상관 {f(G.avg_corr,"{:.2f}")}로 코어 포트(0.19)보다 높다. 같은 테마를 겹쳐 담은 '
  f'결과이며 분산 효과는 {f(G.diversification,"{:.0f}%",pct=True)}에 그친다. 위기에서 상관이 1로 수렴하면 이마저 사라진다.', 9, after=2)
P(f'3. 하락 베타 {f(G.beta_dn,"{:.2f}")}가 상승 베타 {f(G.beta_up,"{:.2f}")}보다 크다. 급락 국면에서 더 민감하다.', 9, after=2)
P('4. 실현 성과 지표(샤프·소르티노·칼마)는 현재 비중을 과거 3년에 고정 적용한 백테스트다. '
  'AI 랠리 구간이 그대로 들어가 선택 편향이 크므로 본문 요약에서 제외했다.', 9, after=2)
P('5. SPCX는 2026년 6월 상장이라 3년 가격 이력이 없다 — 위험 지표 산출에서 제외했다.', 9, after=2)
P('6. 메모리 4사의 낮은 forward P/E는 저평가가 아니라 피크 이익에 붙은 배수일 수 있다. '
  '이 포트 최대 단일 위험이다.', 9, after=2)

H('산출 방법', 12, rule=True)
P('· 유니버스·BM 비중: S&P 500 503종목, 시총 정규화(부동주 미조정)\n'
  '· 성장: 직전 실적 → 차년도 컨센서스까지 2년 CAGR (단년 성장률은 기저효과에 노출)\n'
  '· 목표여력: 애널리스트 목표주가 중앙값 대비\n'
  '· 위험 지표: 3년 일별 수익률, BM 프록시 SPY, 무위험수익률 3.72%(13주 T-bill)\n'
  '· 캡처 비율: 상승·하락일의 평균 수익률 비\n'
  '· 데이터: yfinance(가격·재무·컨센서스), GOOGL 재무는 SEC 8-K 원문 대조 완료\n'
  '· 재현: mp_quant/agent20_{build,chart,docx}.py', 8.5)

P('\n본 자료는 내부 검토용이며 투자 권유가 아니다. 모든 수치는 2026년 8월 26일 기준 실측이다.',
  8, italic=True, color=GREY)

doc.save('AI에이전트_20선_포트폴리오_2026-08-26.docx')
print('생성 완료: AI에이전트_20선_포트폴리오_2026-08-26.docx')
